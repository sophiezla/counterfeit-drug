r"""
Render paper/paper.md into an IEEE Access-formatted, fully editable .docx.

paper.md is the single source of truth; this script converts it rather than
restating it, so the two can never drift. It handles exactly the Markdown
constructs the manuscript uses:

  #/##/###          headings
  pipe tables       -> real Word tables (editable cells, header row bolded)
  > **FIGURE N.**   -> the actual PNG, embedded and centred, plus a caption
  $$ ... $$         -> centred display equation with a right-aligned number
  $ ... $           -> inline maths
  ```...```         -> monospace block
  **bold** *it* `code` <sup>x</sup>   inline runs
  ---               horizontal rule

LaTeX is converted to Unicode by `latex_to_unicode` below. That covers the
constructs this manuscript actually uses; anything unrecognised is passed
through verbatim rather than mangled, so an unsupported command shows up as
readable LaTeX in the output instead of silently disappearing.

------------------------------------------------------------------ IEEE Access
Every metric below is taken from ieeeaccess.cls rather than eyeballed from a
published PDF:

  page          203.2 x 276.2 mm, text block 177.53 mm wide, 672 pt tall
  columns       two, 85.29 mm each, 6.95 mm gutter
  body          Times New Roman 10/12
  title         sans bold 22/25.4          \titlefont
  authors       sans bold 9.9/12           \authorfont
  section A     sans bold 9/12, UPPERCASE, Access blue      \sectionAfont
  section B     sans bold italic 9/12, UPPERCASE            \sectionBfont
  section C     sans regular 9/12                           \sectionCfont
  ABSTRACT /    sans bold 10/12, Access blue, run-in label
  INDEX TERMS
  captions      sans 7/8.4; the "TABLE n."/"FIGURE n." label is bold and blue
  references    Times 7.61/9

Access blue is Pantone 3015 C, approximated here as RGB(0, 84, 143): Word has
no spot-colour model, and IEEE re-typesets accepted papers anyway.

Front matter, the wide tables and the figures are laid out in full-width
one-column bands (the Word equivalent of LaTeX's table*/figure*); prose runs
in two columns. Blocks are emitted in source order, so a band splits the page
where it occurs rather than floating.

Two things in paper.md are internal apparatus and are deliberately NOT
rendered into the manuscript: the "Reference verification status" note, and
the trailing italic verification annotation on each reference entry. They stay
in the source because they are the audit trail; they are not part of a
submission. Set KEEP_INTERNAL_NOTES = True to render them anyway.

Output: paper/PharmaChecked_v2_manuscript.docx
"""
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import (WD_ALIGN_PARAGRAPH, WD_LINE_SPACING,
                            WD_TAB_ALIGNMENT)
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[2]
SRC = ROOT / "paper" / "paper.md"
OUT = ROOT / "paper" / "PharmaChecked_v2_manuscript.docx"

KEEP_INTERNAL_NOTES = False

# ------------------------------------------------------------------ house style
SERIF = "Times New Roman"
SANS = "Arial"                 # metric-compatible stand-in for Helvetica
MONO = "Consolas"

BODY_PT = 10.0
BODY_LEAD = 12.0
CAPTION_PT = 7.0
REF_PT = 7.61

ACCESS_BLUE = RGBColor(0x00, 0x54, 0x8F)     # ~Pantone 3015 C
INK = RGBColor(0x00, 0x00, 0x00)

PAGE_W, PAGE_H = Mm(203.2), Mm(276.2)
MARGIN_X = Mm(12.835)          # (203.2 - 177.53) / 2
MARGIN_TOP, MARGIN_BOT = Mm(19.05), Mm(20.0)
COL_W, COL_SEP = Mm(85.29), Mm(6.95)
TEXT_W = Mm(177.53)


# --------------------------------------------------------------- LaTeX → Unicode
GREEK = {
    "alpha": "α", "beta": "β", "gamma": "γ", "delta": "δ", "Delta": "Δ",
    "epsilon": "ε", "theta": "θ", "lambda": "λ", "mu": "μ", "sigma": "σ",
    "Sigma": "Σ", "tau": "τ", "phi": "φ", "varphi": "φ", "Phi": "Φ",
    "psi": "ψ", "omega": "ω", "Omega": "Ω", "pi": "π", "rho": "ρ", "eta": "η",
}
SYMBOLS = {
    "times": "×", "cdot": "·", "le": "≤", "ge": "≥", "neq": "≠", "approx": "≈",
    "in": "∈", "notin": "∉", "circ": "∘", "star": "*", "pm": "±",
    "rightarrow": "→", "to": "→", "ldots": "…", "dots": "…", "quad": "  ",
    "qquad": "    ", ",": " ", ";": " ", "!": "", " ": " ",
    "sum": "Σ", "prod": "Π", "partial": "∂", "mid": " | ", "|": "‖",
    "infty": "∞", "propto": "∝", "sim": "~", "leq": "≤", "geq": "≥",
}
OPERATORS = ["argmin", "argmax", "min", "max", "log", "exp", "softmax",
             "clip", "resize", "decode", "encode", "diag", "sign"]
SUP = str.maketrans("0123456789+-=()n", "⁰¹²³⁴⁵⁶⁷⁸⁹⁺⁻⁼⁽⁾ⁿ")
SUB = str.maketrans("0123456789+-=()aeoxhklmnpst",
                    "₀₁₂₃₄₅₆₇₈₉₊₋₌₍₎ₐₑₒₓₕₖₗₘₙₚₛₜ")


def _braced(s, i):
    """Return (content, index_after) for a {...} group starting at s[i]=='{'."""
    assert s[i] == "{"
    depth, j = 0, i
    while j < len(s):
        if s[j] == "{":
            depth += 1
        elif s[j] == "}":
            depth -= 1
            if depth == 0:
                return s[i + 1:j], j + 1
        j += 1
    return s[i + 1:], len(s)


def latex_to_unicode(s: str) -> str:
    s = s.strip()
    s = s.replace(r"\{", "\x01").replace(r"\}", "\x02")
    s = s.replace(r"\begin{cases}", "\x01 ").replace(r"\end{cases}", " \x02")
    s = s.replace("\\\\", "\x03")
    s = re.sub(r"\\(?:left|right|big|Big|bigg|Bigg)\s*", "", s)
    s = re.sub(r"\\mathcal\{L\}", "\u2112", s)
    s = re.sub(r"\\mathcal\{([A-Za-z])\}", r"\1", s)
    for op in OPERATORS:
        s = s.replace("\\" + op, op)
    s = re.sub(r"\\mathrm\{([^{}]*)\}", r"\1", s)
    s = re.sub(r"\\mathrm\s*", "", s)
    s = re.sub(r"\\text\{([^{}]*)\}", r"\1", s)
    s = re.sub(r"\\operatorname\{([^{}]*)\}", r"\1", s)
    s = s.replace(r"\mathbb{R}", "ℝ").replace(r"\mathbb{1}", "1")
    s = re.sub(r"\\mathbb\{E\}", "E", s)
    s = re.sub(r"\\mathbf\{([^{}]*)\}", r"\1", s)

    def frac(match_src):
        out, i = "", 0
        while i < len(match_src):
            m = re.compile(r"\\[d]?frac").match(match_src, i)
            if m and match_src[m.end():m.end() + 1] == "{":
                num, j = _braced(match_src, m.end())
                if j < len(match_src) and match_src[j] == "{":
                    den, k = _braced(match_src, j)
                    out += f"({frac(num)})/({frac(den)})"
                    i = k
                    continue
            out += match_src[i]
            i += 1
        return out
    s = frac(s)

    s = re.sub(r"\\sqrt\{([^{}]*)\}", r"√(\1)", s)

    def binom(text):
        out, i = "", 0
        while i < len(text):
            m = re.compile(r"\\binom").match(text, i)
            if m and text[m.end():m.end() + 1] == "{":
                a, j = _braced(text, m.end())
                if j < len(text) and text[j] == "{":
                    b, k = _braced(text, j)
                    out += f"C({binom(a)}, {binom(b)})"
                    i = k
                    continue
            out += text[i]
            i += 1
        return out
    s = binom(s)
    s = re.sub(r"\\hat\{([A-Za-z])\}", lambda m: m.group(1) + "\u0302", s)
    s = re.sub(r"\\bar\{([A-Za-z])\}", lambda m: m.group(1) + "\u0304", s)
    s = re.sub(r"\\overline\{([^{}]*)\}", lambda m: m.group(1) + "\u0304", s)
    s = s.replace(r"\top", "\u1d40")

    for name, ch in sorted(GREEK.items(), key=lambda kv: -len(kv[0])):
        s = s.replace("\\" + name, ch)
    for name, ch in sorted(SYMBOLS.items(), key=lambda kv: -len(kv[0])):
        s = s.replace("\\" + name, ch)

    def script(text, marker, table):
        out, i = "", 0
        while i < len(text):
            if text[i] == marker and i + 1 < len(text):
                if text[i + 1] == "{":
                    grp, j = _braced(text, i + 1)
                    if all(ord(c) in table for c in grp):
                        out += grp.translate(table)
                    else:
                        out += marker + "(" + grp + ")"
                    i = j
                    continue
                ch = text[i + 1]
                out += ch.translate(table) if ord(ch) in table else marker + ch
                i += 2
                continue
            out += text[i]
            i += 1
        return out
    s = script(s, "^", SUP)
    s = script(s, "_", SUB)

    s = s.replace("\x03", " ;  ").replace("&", " ")
    s = re.sub(r"[{}]", "", s)
    s = s.replace("\x01", "{").replace("\x02", "}")
    s = re.sub(r"\s{3,}", "   ", s).strip()
    if "\\" in s:
        print(f"    NOTE unconverted LaTeX in: {s}")
    return s


# --------------------------------------------------------------- docx helpers
def set_cell_background(cell, hex_colour):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_colour)
    tcPr.append(shd)


def set_columns(section, count):
    """IEEE Access is a two-column journal; full-width bands are one column."""
    sectPr = section._sectPr
    for old in sectPr.findall(qn("w:cols")):
        sectPr.remove(old)
    cols = OxmlElement("w:cols")
    cols.set(qn("w:num"), str(count))
    cols.set(qn("w:equalWidth"), "1")
    cols.set(qn("w:space"), str(int(COL_SEP.twips)) if count > 1 else "0")
    cols.set(qn("w:sep"), "0")
    sectPr.append(cols)


def page_setup(section):
    section.page_width, section.page_height = PAGE_W, PAGE_H
    section.left_margin = section.right_margin = MARGIN_X
    section.top_margin, section.bottom_margin = MARGIN_TOP, MARGIN_BOT


def band(doc, count):
    """Open a new continuous section with `count` columns."""
    s = doc.add_section(WD_SECTION.CONTINUOUS)
    page_setup(s)
    set_columns(s, count)
    return s


def add_page_number_field(paragraph):
    for instr, kind in (("begin", "w:fldChar"), ("PAGE", "w:instrText"),
                        ("end", "w:fldChar")):
        el = OxmlElement(kind)
        if kind == "w:fldChar":
            el.set(qn("w:fldCharType"), instr)
        else:
            el.set(qn("xml:space"), "preserve")
            el.text = " PAGE "
        run = paragraph.add_run()
        run.font.name = SANS
        run.font.size = Pt(8)
        run._r.append(el)


def build_footer(section):
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("VOLUME 14, 2026   ")
    r.font.name, r.font.size, r.font.color.rgb = SANS, Pt(8), ACCESS_BLUE
    r.bold = True
    add_page_number_field(p)


INLINE = re.compile(
    r"(\*\*.+?\*\*|\*[^*]+?\*|`[^`]+?`|\$[^$]+?\$|<sup>.*?</sup>)")


def add_runs(paragraph, text, size=None, font=None, colour=None):
    """Render inline markdown + inline maths into runs on `paragraph`."""
    made = []
    for part in INLINE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            r = paragraph.add_run(part[2:-2]); r.bold = True
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            r = paragraph.add_run(part[1:-1]); r.italic = True
        elif part.startswith("<sup>"):
            r = paragraph.add_run(re.sub(r"</?sup>", "", part))
            r.font.superscript = True
        elif part.startswith("`") and part.endswith("`"):
            r = paragraph.add_run(part[1:-1])
            r.font.name = MONO
            r.font.size = Pt((size or BODY_PT) - 1.0)
        elif part.startswith("$") and part.endswith("$"):
            r = paragraph.add_run(latex_to_unicode(part[1:-1]))
            r.italic = True
        else:
            r = paragraph.add_run(part)
        made.append(r)
    for r in made:
        if font and r.font.name != MONO:
            r.font.name = font
        if size and r.font.name != MONO:
            r.font.size = Pt(size)
        if colour:
            r.font.color.rgb = colour
    return made


def body_paragraph(doc, text, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                   space_after=0, size=BODY_PT, font=SERIF, indent=None,
                   lead=BODY_LEAD):
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(0)
    pf.line_spacing = Pt(lead)
    if indent is not None:
        pf.first_line_indent = Inches(indent)
    add_runs(p, text, size=size, font=font)
    return p


# --------------------------------------------------------------- markdown parse
def parse_blocks(md):
    lines = md.split("\n")
    blocks, i = [], 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        if stripped.startswith("```"):
            j = i + 1
            buf = []
            while j < len(lines) and not lines[j].strip().startswith("```"):
                buf.append(lines[j]); j += 1
            blocks.append(("code", buf))
            i = j + 1
            continue

        if stripped.startswith("$$"):
            if stripped.endswith("$$") and len(stripped) > 4:
                blocks.append(("math", stripped))
                i += 1
                continue
            buf, j = [stripped], i + 1
            while j < len(lines):
                buf.append(lines[j].strip())
                if lines[j].strip().endswith("$$"):
                    break
                j += 1
            blocks.append(("math", " ".join(buf)))
            i = j + 1
            continue

        if stripped == "---":
            blocks.append(("rule", None)); i += 1; continue

        m = re.match(r"^(#{1,4})\s+(.*)$", stripped)
        if m:
            blocks.append((f"h{len(m.group(1))}", m.group(2))); i += 1; continue

        if stripped.startswith(">"):
            buf = []
            while i < len(lines) and lines[i].strip().startswith(">"):
                buf.append(lines[i].strip().lstrip(">").strip()); i += 1
            blocks.append(("quote", buf))
            continue

        if stripped.startswith("|"):
            buf = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                buf.append(lines[i].strip()); i += 1
            blocks.append(("table", buf))
            continue

        if re.match(r"^(\d+\.|[-*])\s+", stripped):
            buf = []
            while i < len(lines) and (
                    re.match(r"^(\d+\.|[-*])\s+", lines[i].strip())
                    or (lines[i].startswith("   ") and lines[i].strip())):
                if re.match(r"^(\d+\.|[-*])\s+", lines[i].strip()):
                    buf.append(re.sub(r"^(\d+\.|[-*])\s+", "", lines[i].strip()))
                else:
                    buf[-1] += " " + lines[i].strip()
                i += 1
            ordered = bool(re.match(r"^\d+\.", stripped))
            blocks.append(("list", (ordered, buf)))
            continue

        buf = [stripped]
        i += 1
        while i < len(lines) and lines[i].strip() and not re.match(
                r"^(#{1,4}\s|\||>|```|\$\$|---$|\d+\.\s|[-*]\s)", lines[i].strip()):
            buf.append(lines[i].strip()); i += 1
        blocks.append(("para", " ".join(buf)))
    return blocks


# --------------------------------------------------------------- renderers
FIG_RE = re.compile(
    r"^\*\*(FIGURE\s+\d+\.)\*\*\s*`([^`]+)`\s*[—-]\s*(.*)$", re.S)
TABLE_CAP_RE = re.compile(r"^\*\*(TABLE\s+\d+\.)\*\*\s*(.*)$", re.S)
EQ_NUM_RE = re.compile(r"\\tag\{(\d+)\}")
REF_RE = re.compile(r"^\[\d+\]\s")
# A paragraph opening with an escaped asterisk, directly after a table, is
# that table's tabular note. build_tex.py folds it into the float; here
# floats are inline, so it only needs a smaller size than body text.
TABLE_NOTE_RE = re.compile(r"^\\\*\s+")
# A reference entry may carry a trailing italic verification annotation: the
# project's audit trail, never part of the citation. It is stripped from both
# built artefacts unless KEEP_INTERNAL_NOTES is set. The rule is structural
# rather than a whitelist of opening words -- an earlier whitelist silently
# passed six notes through into the compiled PDF.
REF_NOTE_RE = re.compile(r"\s*\*([^*]+)\*\s*$")


def strip_ref_note(entry):
    """Drop a reference entry's trailing internal annotation, if it has one.

    A note is a trailing italic span of at least four words that ends in a
    full stop. A bibliographic italic at the end of an entry -- a journal or
    book title -- is shorter than that and carries no terminal period inside
    the italics. Kept identical to build_tex.strip_ref_note; the test is in
    Python because the obvious regex for it backtracks catastrophically on a
    long note.
    """
    m = REF_NOTE_RE.search(entry)
    if not m:
        return entry
    # A note may close with a quotation or a parenthesis after its full stop.
    note = m.group(1).rstrip("\"')]}”’")
    if note.endswith(".") and len(m.group(1).split()) >= 4:
        return entry[:m.start()]
    return entry


def caption_paragraph(doc, label, text, above):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(2 if above else 4)
    pf.space_after = Pt(4 if above else 8)
    pf.line_spacing = Pt(8.4)
    r = p.add_run(label + " ")
    r.bold = True
    r.font.name, r.font.size, r.font.color.rgb = SANS, Pt(CAPTION_PT), ACCESS_BLUE
    add_runs(p, text, size=CAPTION_PT, font=SANS, colour=INK)
    return p


def render_figure(doc, m, full_width):
    label, relpath, caption = m.groups()
    png = ROOT / relpath.replace(".pdf", ".png")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True     # keep the caption attached
    # The Normal style sets EXACT 12 pt line spacing, which is what IEEE
    # Access's 10/12 body requires -- and Word clips an inline image to the
    # exact line height, which silently reduced every figure in this
    # manuscript to a 12 pt sliver of its own bottom edge. The image
    # paragraph must therefore opt back out to single spacing.
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    width = Mm(170) if full_width else Mm(82)
    if png.exists():
        p.add_run().add_picture(str(png), width=width)
    else:
        r = p.add_run(f"[missing figure: {png}]")
        r.italic = True
    caption_paragraph(doc, label, caption, above=False)


def render_quote(doc, lines, full_width):
    """A blockquote holds either a note or one figure per line."""
    pending = []
    for line in lines + [None]:
        if line is not None and FIG_RE.match(line):
            if pending:
                body_paragraph(doc, " ".join(pending), size=BODY_PT - 1.0,
                               space_after=4)
                pending = []
            render_figure(doc, FIG_RE.match(line), full_width)
        elif line:
            pending.append(line)
    if pending:
        body_paragraph(doc, " ".join(pending), size=BODY_PT - 1.0, space_after=4)


def render_math(doc, raw):
    inner = raw.strip().strip("$").strip()
    tag = EQ_NUM_RE.search(inner)
    number = tag.group(1) if tag else None
    inner = EQ_NUM_RE.sub("", inner)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(latex_to_unicode(inner))
    r.italic = True
    r.font.name, r.font.size = SERIF, Pt(BODY_PT)
    if number:
        p.paragraph_format.tab_stops.add_tab_stop(COL_W, WD_TAB_ALIGNMENT.RIGHT)
        r = p.add_run("\t(" + number + ")")
        r.font.name, r.font.size = SERIF, Pt(BODY_PT)


def table_rows(lines):
    rows = []
    for ln in lines:
        if re.match(r"^\|[\s:\-\|]+\|$", ln):
            continue
        rows.append([c.strip() for c in ln.strip().strip("|").split("|")])
    return rows


def render_table(doc, lines, full_width):
    rows = table_rows(lines)
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    rows = [r + [""] * (ncols - len(r)) for r in rows]
    t = doc.add_table(rows=len(rows), cols=ncols)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = True
    size = CAPTION_PT if ncols <= 8 else CAPTION_PT - 0.5
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.cell(ri, ci)
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.space_before = Pt(1)
            p.alignment = (WD_ALIGN_PARAGRAPH.LEFT if ci == 0
                           else WD_ALIGN_PARAGRAPH.CENTER)
            add_runs(p, val, size=size, font=SANS)
            if ri == 0:
                for run in p.runs:
                    run.bold = True
                set_cell_background(cell, "DCE6EF")
    doc.add_paragraph().paragraph_format.space_after = Pt(6)


def heading(doc, text, level):
    """level 1 = 'I. INTRODUCTION', 2 = 'A. SUBSECTION', 3 = run-in."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(10 if level == 1 else 8)
    pf.space_after = Pt(3)
    pf.line_spacing = Pt(12)
    pf.keep_with_next = True
    runs = add_runs(p, text.upper() if level <= 2 else text,
                    size=9.0, font=SANS)
    for r in runs:
        r.bold = level <= 2
        r.italic = level == 2
        r.font.color.rgb = ACCESS_BLUE if level == 1 else INK
    return p


# --------------------------------------------------------------- front matter
def render_front_matter(doc, md):
    """Title block, ABSTRACT and INDEX TERMS, full width, before the columns."""
    head = md.split("## I. Introduction")[0]
    for raw in head.split("\n"):
        line = raw.strip()
        if not line or line == "---":
            continue

        if line.startswith("# "):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(10)
            p.paragraph_format.line_spacing = Pt(25.4)
            for r in add_runs(p, line[2:], size=22.0, font=SANS):
                r.bold = True
                r.font.color.rgb = INK
            continue

        if line.startswith("**ABSTRACT**") or line.startswith("**INDEX TERMS**"):
            label, rest = line.split("**", 2)[1], line.split("**", 2)[2]
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.space_after = Pt(8)
            p.paragraph_format.line_spacing = Pt(BODY_LEAD)
            r = p.add_run(label + " ")
            r.bold = True
            r.font.name, r.font.size, r.font.color.rgb = (
                SANS, Pt(10), ACCESS_BLUE)
            add_runs(p, rest.strip(), size=BODY_PT, font=SERIF)
            continue

        # author line, affiliations, corresponding author, funding note
        is_authors = line.startswith("**[AUTHOR")
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(6 if is_authors else 1)
        p.paragraph_format.line_spacing = Pt(12)
        for r in add_runs(p, line, size=9.9 if is_authors else 8.0,
                          font=SANS if is_authors else SERIF):
            if not is_authors:
                r.italic = True


# --------------------------------------------------------------- main
def main():
    md = SRC.read_text(encoding="utf-8")
    blocks = parse_blocks(md)
    # everything from the Introduction onward is rendered from blocks; the
    # front matter is rendered from raw lines so affiliations keep their breaks
    start = next(i for i, (k, v) in enumerate(blocks)
                 if k == "h2" and str(v).startswith("I. Introduction"))
    blocks = blocks[start:]

    doc = Document()
    st = doc.styles["Normal"]
    st.font.name = SERIF
    st.font.size = Pt(BODY_PT)
    st._element.rPr.rFonts.set(qn("w:eastAsia"), SERIF)
    st.paragraph_format.space_after = Pt(0)
    st.paragraph_format.line_spacing = Pt(BODY_LEAD)

    sec = doc.sections[0]
    page_setup(sec)
    set_columns(sec, 1)
    build_footer(sec)

    render_front_matter(doc, md)
    band(doc, 2)

    cols = 2
    in_references = False
    n_bands = 0

    def want(count):
        nonlocal cols, n_bands
        if cols != count:
            band(doc, count)
            cols = count
            n_bands += 1

    for idx, (kind, payload) in enumerate(blocks):
        nxt = blocks[idx + 1] if idx + 1 < len(blocks) else (None, None)

        # tables, their captions and figures go in full-width bands
        is_table_caption = (kind == "para" and TABLE_CAP_RE.match(str(payload))
                            and nxt[0] == "table")
        is_figure_quote = (kind == "quote"
                           and any(FIG_RE.match(l) for l in payload))
        if is_table_caption or kind == "table" or is_figure_quote:
            want(1)
        elif kind in ("h2", "h3", "h4", "para", "list", "math", "code", "quote"):
            want(2)

        if kind == "h2":
            text = str(payload)
            in_references = text.startswith("References")
            heading(doc, text, level=1)
        elif kind == "h3":
            heading(doc, str(payload), level=2)
        elif kind == "h4":
            heading(doc, str(payload), level=3)
        elif kind == "para":
            text = str(payload)
            img = re.match(r"^!\[[^\]]*\]\(([^)]+)\)$", text)
            if img:                       # the author photograph
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                p.paragraph_format.space_after = Pt(4)
                path = ROOT / img.group(1)
                if path.exists():
                    p.add_run().add_picture(str(path), width=Inches(1.0))
                else:
                    p.add_run(f"[missing image: {path}]").italic = True
                continue
            m = TABLE_CAP_RE.match(text)
            if m:
                caption_paragraph(doc, m.group(1), m.group(2), above=True)
            elif TABLE_NOTE_RE.match(text) and idx and blocks[idx - 1][0] == "table":
                # A tabular note: it defines a mark used in the cells above.
                # Floats are inline here, so it already sits under its table
                # and only needs the smaller size that marks it as apparatus
                # rather than body text -- and the escaped asterisk unescaped.
                body_paragraph(doc, TABLE_NOTE_RE.sub("* ", text),
                               size=REF_PT, space_after=6, lead=9.0, indent=0)
            elif in_references and REF_RE.match(text):
                if not KEEP_INTERNAL_NOTES:
                    text = strip_ref_note(text)
                p = body_paragraph(doc, text, size=REF_PT, space_after=3,
                                   lead=9.0)
                p.paragraph_format.left_indent = Mm(5)
                p.paragraph_format.first_line_indent = Mm(-5)
            else:
                body_paragraph(doc, text, space_after=0, indent=0.14)
        elif kind == "quote":
            if (not KEEP_INTERNAL_NOTES and payload
                    and payload[0].startswith("**Reference verification")):
                continue
            render_quote(doc, payload, full_width=(cols == 1))
        elif kind == "math":
            render_math(doc, payload)
        elif kind == "table":
            render_table(doc, payload, full_width=(cols == 1))
        elif kind == "list":
            ordered, items = payload
            for it in items:
                p = doc.add_paragraph(
                    style="List Number" if ordered else "List Bullet")
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.line_spacing = Pt(BODY_LEAD)
                add_runs(p, it, size=BODY_PT, font=SERIF)
        elif kind == "code":
            for ln in payload:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.line_spacing = Pt(8.5)
                p.paragraph_format.left_indent = Mm(3)
                r = p.add_run(ln if ln else " ")
                r.font.name = MONO
                r.font.size = Pt(7.0)
            doc.add_paragraph().paragraph_format.space_after = Pt(4)
        elif kind == "rule":
            continue

    doc.save(OUT)
    try:                                   # OUT is patched when Word holds the
        shown = OUT.relative_to(ROOT)      # canonical file and we build to temp
    except ValueError:
        shown = OUT
    print(f"wrote {shown}")
    print(f"  {len(blocks)} blocks, {len(doc.tables)} tables, "
          f"{len(doc.inline_shapes)} embedded figures, "
          f"{n_bands} column bands, {len(doc.sections)} sections")


if __name__ == "__main__":
    main()
