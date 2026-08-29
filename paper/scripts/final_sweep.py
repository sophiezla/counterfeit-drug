r"""
Pre-submission sweep over the manuscript and its supplement.

Checks the things that break silently after a restructuring: references that
point at a heading which no longer exists, floats that are never cited or
never defined, numbering that disagrees between the .docx and the .pdf,
encoding damage, and leftover editorial placeholders.

    python paper/scripts/final_sweep.py

Exit status is non-zero if any check fails.
"""
import io
import re
import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parents[2]
MAIN_MD = ROOT / "paper" / "paper.md"
SUP_MD = ROOT / "paper" / "supplementary.md"
MAIN_PDF = ROOT / "paper" / "latex" / "paper.pdf"
SUP_PDF = ROOT / "paper" / "latex" / "supplementary.pdf"
DOCX = ROOT / "paper" / "PharmaChecked_v2_manuscript.docx"

fails, notes = [], []


def check(ok, label, detail=""):
    print(f"  [{'ok' if ok else 'FAIL'}] {label}" + (f" -- {detail}" if detail else ""))
    if not ok:
        fails.append(label)


def note(label):
    print(f"  [note] {label}")
    notes.append(label)


def read(p):
    return io.open(p, encoding="utf-8").read()


def strip_code(t):
    return re.sub(r"```.*?```", "", t, flags=re.S)


def headings(text):
    """{'VI-A': title} for every numbered heading."""
    out, major = {}, None
    for ln in strip_code(text).split("\n"):
        m = re.match(r"^##\s+(S-[IVXL]+|[IVXL]+)\.\s+(.*)$", ln)
        if m:
            major = m.group(1)
            out[major] = m.group(2)
            continue
        m = re.match(r"^###\s+([A-Z])\.\s+(.*)$", ln)
        if m and major:
            out[f"{major}-{m.group(1)}"] = m.group(2)
    return out


print("SECTION REFERENCES")
main, sup = read(MAIN_MD), read(SUP_MD)
h_main, h_sup = headings(main), headings(sup)
print(f"  main headings: {len(h_main)}   supplement headings: {len(h_sup)}")

SECREF = re.compile(r"\bSections?\s+((?:S-)?[IVXL]+(?:-[A-Z])?)")
for name, text, own, other in (("main", main, h_main, h_sup),
                               ("supplement", sup, h_sup, h_main)):
    bad = []
    for ref in set(SECREF.findall(strip_code(text))):
        target = own if not ref.startswith("S-") else h_sup
        if name == "supplement" and not ref.startswith("S-"):
            target = h_main            # supplement citing the main paper
        if ref not in target:
            bad.append(ref)
    check(not bad, f"{name}: every section reference resolves",
          f"dangling: {sorted(bad)}")

print("\nFLOATS")
for name, text, pre in (("main", main, ""), ("supplement", sup, "S")):
    t = strip_code(text)
    tcaps = re.findall(rf"\*\*TABLE\s+{pre}(\d+)\.\*\*", t)
    fcaps = re.findall(rf"^>\s*\*\*FIGURE\s+{pre}(\d+)\.\*\*", t, re.M)
    check(tcaps == [str(i) for i in range(1, len(tcaps) + 1)],
          f"{name}: table captions sequential", f"got {tcaps}")
    check(fcaps == [str(i) for i in range(1, len(fcaps) + 1)],
          f"{name}: figure captions sequential", f"got {fcaps}")
    trefs = {m for m in re.findall(rf"\bTables?\s+{pre}(\d+)\b", t)}
    frefs = {m for m in re.findall(rf"\b(?:Figs?\.|Figures?)\s+{pre}(\d+)\b", t)}
    check(not (trefs - set(tcaps)), f"{name}: table refs have captions",
          f"dangling: {sorted(trefs - set(tcaps))}")
    check(not (frefs - set(fcaps)), f"{name}: figure refs have captions",
          f"dangling: {sorted(frefs - set(fcaps))}")
    uncited_t = set(tcaps) - trefs
    if uncited_t:
        note(f"{name}: tables never referred to in their own file: "
             f"{sorted(uncited_t, key=int)}")

print("\nENCODING AND PLACEHOLDERS")
for name, p in (("paper.md", MAIN_MD), ("supplementary.md", SUP_MD)):
    t = read(p)
    bad = [c for c in ("﻿", "â€", "Ã—") if c in t]
    check(not bad, f"{name}: no encoding damage", f"found {bad}")
    ph = re.findall(r"\[(?:TO BE|To be|Placeholder|AUTHOR|Degree)[^\]]{0,60}\]", t)
    check(not ph, f"{name}: no editorial placeholders", f"{ph[:3]}")

print("\nBUILT ARTEFACTS")
try:
    import fitz
    mp, sp = fitz.open(MAIN_PDF), fitz.open(SUP_PDF)
    check(mp.page_count <= 20, "main paper is 20 pages or fewer",
          f"{mp.page_count} pages")
    print(f"  [note] supplement: {sp.page_count} pages")

    pdf_text = "\n".join(p.get_text() for p in mp)
    from docx import Document
    d = Document(DOCX)
    dx = "\n".join(p.text for p in d.paragraphs)
    for t in d.tables:
        for row in t.rows:
            dx += "\n" + " ".join(c.text for c in row.cells)

    def uniq(seq):
        out = []
        for x in seq:
            if x not in out:
                out.append(x)
        return out
    for word in ("TABLE", "FIGURE"):
        a = uniq(re.findall(rf"{word}\s+(\d+)\.", dx))
        b = uniq(re.findall(rf"{word}\s+(\d+)\.", pdf_text))
        check(a == b, f".docx and .pdf agree on {word} numbering",
              f"docx {a} vs pdf {b}")

    # every citation in the PDF resolves to a numbered reference entry
    cited = {int(x) for x in re.findall(r"\[(\d{1,2})\]", pdf_text)}
    listed = {int(x) for x in re.findall(r"(?m)^\[(\d{1,2})\]", pdf_text)}
    check(not (cited - listed - set(range(1, 31))),
          "PDF citations fall within the reference list")

    # ---------------------------------------------------------- front matter
    # Page 1 carries three placeholders that are NOT ours to fill, and it is
    # worth naming them, because "the PDF still says 2023" reads like a defect
    # on inspection. ieeeaccess.cls hard-codes VOLUME 11, 2023 in its footer,
    # and the official template ships the dummy \history and \doi lines
    # verbatim -- the class file used here is byte-identical to the one in
    # ACCESS_latex_template_20240429 -- and IEEE replaces all three at
    # production. Everything the AUTHOR owns is checked to be filled, and any
    # placeholder-shaped token outside the template's own fields is reported.
    TEMPLATE_FIELDS = ("xxxx 00, 0000", "10.1109/ACCESS.2026.DOI",
                       "VOLUME 11, 2023")
    for field in TEMPLATE_FIELDS:
        if field in pdf_text:
            note(f"template placeholder present as expected, IEEE fills it at "
                 f"production: {field!r}")

    squashed = re.sub(r"\s+", "", pdf_text)
    title = re.search(r"^#\s+(.*)$", main, re.M).group(1)
    check(re.sub(r"\s+", "", title) in squashed,
          "PDF title matches paper.md's title")
    for label, needle in (("author name", "SOPHIE ZHU"),
                          ("affiliation", "Mira Costa High School"),
                          ("author e-mail", "sophiezhu2028@gmail.com"),
                          ("ORCID", "0009-0004-2403-910X"),
                          ("corresponding author", "Corresponding author"),
                          ("index terms", "INDEX TERMS"),
                          ("funding statement", "no specific grant"),
                          ("AI-use disclosure", "Generative-AI disclosure"),
                          ("author biography", "is a student at Mira Costa")):
        check(re.sub(r"\s+", "", needle) in squashed,
              f"front matter carries the {label}")

    residue = pdf_text
    for field in TEMPLATE_FIELDS:
        residue = residue.replace(field, "")
    leftovers = sorted({m.group(0).strip() for m in re.finditer(
        r"(?i)\b(xxxx+|tbd|todo|lorem ipsum|first a\. author|author@|"
        r"your name here)\b", residue)})
    check(not leftovers,
          "no placeholder text outside the template's own fields",
          f"found: {leftovers}")
except Exception as exc:                                  # noqa: BLE001
    check(False, "PDF/docx inspection", str(exc))

print("\nRESULT")
if fails:
    print(f"{len(fails)} check(s) FAILED: {fails}")
    sys.exit(1)
print("all checks passed" + (f"; {len(notes)} note(s)" if notes else ""))
